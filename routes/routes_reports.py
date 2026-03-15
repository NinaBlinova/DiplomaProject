import io
from collections import defaultdict

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from flask import request, send_file, Blueprint
from docx import Document
from docx.shared import Inches

from model.data.LoggerService import LoggerService
from model.database import DatabaseEngine


report_bp = Blueprint("report_bp", __name__, url_prefix="/api")

db_engine = DatabaseEngine()
logger = LoggerService(db_engine)


def format_number(value):
    try:
        return f"{value:,.2f}"
    except Exception:
        return "0"


def group_data_by_year(data):
    years = defaultdict(list)
    for row in data:
        years[row.get("Year")].append(row)
    return years


def add_table(document, rows):
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    header = table.rows[0].cells
    header[0].text = "Номер месяца"
    header[1].text = "Доход (руб.)"
    header[2].text = "Налог (руб.)"
    header[3].text = "Количество транзакций"

    rows = sorted(rows, key=lambda x: x.get("Month", 0))

    months = []
    incomes = []
    taxes = []
    transactions = []

    for row in rows:
        month = row.get("Month", 0)
        income = row.get("Income", 0)
        tax = row.get("Tax", 0)
        trans = row.get("Transactions", 0)

        cells = table.add_row().cells
        cells[0].text = str(month)
        cells[1].text = format_number(income)
        cells[2].text = format_number(tax)
        cells[3].text = str(trans)

        months.append(month)
        incomes.append(income)
        taxes.append(tax)
        transactions.append(trans)

    return months, incomes, taxes, transactions


def create_plot(months, incomes, taxes, transactions, year):
    plt.figure()

    plt.plot(months, incomes, label="Доход (руб.)")
    plt.plot(months, taxes, label="Налог (руб.)")
    plt.plot(months, transactions, label="Количество транзакций")

    plt.title(f"Статистика за {year} год")
    plt.xlabel("Месяц")
    plt.legend()

    img_stream = io.BytesIO()
    plt.savefig(img_stream)
    plt.close()

    img_stream.seek(0)
    return img_stream


@report_bp.route("/report", methods=["POST"])
def generate_report():
    data = request.json

    if not data:
        return {"error": "No data provided"}, 400

    user = data.get("user", {})
    median_data = data.get("medianData", [])

    if not median_data:
        return {"error": "medianData is empty"}, 400

    filters = data.get("filters", {})

    user_id = user.get("Id")
    username = user.get("Username", "Unknown")
    user_fullname = user.get("FullName", "Не указано")

    model_name = median_data[0].get("ModelName")
    model_version = median_data[0].get("ModelVersion")

    tax_type = filters.get("taxType", "Не указано")
    inn = filters.get("inn", "Не указано")

    try:
        document = Document()

        document.add_heading("Отчет", level=0)
        document.add_paragraph(f"Отчет создан пользователем: {user_fullname}")
        document.add_paragraph(f"Модель: {model_name}")
        document.add_paragraph(f"Версия модели: {model_version}")
        document.add_paragraph(f"Тип налога: {tax_type}")
        document.add_paragraph(f"ИНН: {inn}")

        years = group_data_by_year(median_data)

        for year, rows in sorted(years.items()):
            document.add_heading(f"Год {year}", level=1)

            months, incomes, taxes, transactions = add_table(document, rows)

            document.add_paragraph("")

            img_stream = create_plot(months, incomes, taxes, transactions, year)
            document.add_picture(img_stream, width=Inches(6))

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        logger.log_action(
            user_id=user_id,
            username=username,
            action="Сгенерирован отчет",
            additional_info=f"Пользователь {username} сгенерировал отчет. "
                            f"Модель {model_name} версии {model_version}."
        )

        return send_file(
            buffer,
            as_attachment=True,
            download_name="report.docx"
        )

    except Exception as e:

        logger.log_action(
            user_id=user_id,
            username=username,
            action="Ошибка генерации отчета",
            additional_info=str(e)
        )

        return {"error": "Ошибка генерации отчета"}, 500
